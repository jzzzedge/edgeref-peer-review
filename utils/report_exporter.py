"""Report export module — supports Markdown and Word format export."""
from __future__ import annotations

from datetime import datetime
from io import BytesIO
from typing import Dict, Optional

from docx import Document


def build_report_metadata(
    mode: str,
    engine: str,
    journal_profile: Optional[Dict[str, str]] = None,
    manuscript_language: Optional[str] = None,
    review_output_language: Optional[str] = None,
    journal_match: Optional[Dict] = None,
) -> Dict[str, str]:
    """Build report metadata. Does NOT include API keys, provider, or model."""
    now = datetime.now()
    result = {
        "title": "EdgeRef AI Pre-submission Review Report",
        "generated_at": now.strftime("%Y-%m-%d %H:%M:%S"),
        "timestamp": now.strftime("%Y%m%d_%H%M%S"),
        "mode": mode,
        "engine": engine,
    }
    if manuscript_language:
        result["manuscript_language"] = manuscript_language
    if review_output_language:
        result["review_output_language"] = review_output_language
    if journal_profile:
        result["target_journal"] = journal_profile.get("journal_name", "")
        result["subject_area"] = journal_profile.get("subject_area", "")
        result["article_type"] = journal_profile.get("article_type", "")
    if journal_match:
        if journal_match.get("match_status") == "matched":
            result["journal_match_status"] = "Matched"
            result["journal_match_confidence"] = journal_match.get("confidence", "")
            result["journal_match_journal"] = journal_match.get("matched_journal_name", "")
            result["journal_match_database"] = journal_match.get("database_summary") or journal_match.get("database", "")
            result["journal_match_count"] = str(journal_match.get("match_count", ""))
            result["journal_match_db_type"] = journal_match.get("database_type", "")
            result["journal_match_version"] = journal_match.get("version", "")
            result["journal_match_subject"] = journal_match.get("subject_area", "")
            result["journal_match_issn"] = journal_match.get("issn", "")
            result["journal_match_cn"] = journal_match.get("cn", "")
        elif journal_match.get("match_status") == "not_matched":
            result["journal_match_status"] = "Not matched"
        else:
            result["journal_match_status"] = "Not checked"
    return result


def export_markdown_report(report_text: str, metadata: Dict[str, str]) -> bytes:
    """Export report as Markdown, returns bytes."""
    if not report_text or not report_text.strip():
        return "_(Report content is empty.)_".encode("utf-8")
    lines = [
        f"# {metadata.get('title', 'EdgeRef AI Pre-submission Review Report')}",
        "",
        f"- **Generated**: {metadata.get('generated_at', '')}",
        f"- **Review Mode**: {metadata.get('mode', '')}",
        f"- **Review Engine**: {metadata.get('engine', '')}",
    ]

    if metadata.get("manuscript_language"):
        lines.append(f"- **Manuscript Language**: {metadata['manuscript_language']}")
    if metadata.get("review_output_language"):
        lines.append(f"- **Review Output Language**: {metadata['review_output_language']}")
    if metadata.get("journal_match_status"):
        lines.append(f"- **Journal Catalog Match**: {metadata['journal_match_status']}")
        if metadata.get("journal_match_confidence"):
            lines.append(f"  - Confidence: {metadata['journal_match_confidence']}")
        if metadata.get("journal_match_count"):
            lines.append(f"  - Matched records: {metadata['journal_match_count']}")
        if metadata.get("journal_match_journal"):
            lines.append(f"  - Matched Journal: {metadata['journal_match_journal']}")
        if metadata.get("journal_match_database"):
            lines.append(f"  - Database: {metadata['journal_match_database']}")
        if metadata.get("journal_match_db_type"):
            lines.append(f"  - DB Type: {metadata['journal_match_db_type']}")
        if metadata.get("journal_match_version"):
            lines.append(f"  - Version: {metadata['journal_match_version']}")
        if metadata.get("journal_match_issn"):
            lines.append(f"  - ISSN: {metadata['journal_match_issn']}")
        if metadata.get("journal_match_cn"):
            lines.append(f"  - CN: {metadata['journal_match_cn']}")

    if metadata.get("target_journal"):
        lines.append(f"- **Target Journal**: {metadata['target_journal']}")
    if metadata.get("subject_area"):
        lines.append(f"- **Subject Area**: {metadata['subject_area']}")
    if metadata.get("article_type"):
        lines.append(f"- **Article Type**: {metadata['article_type']}")

    lines.extend(["", "---", "", report_text or ""])
    return "\n".join(lines).encode("utf-8")


def export_docx_report(report_text: str, metadata: Dict[str, str]) -> BytesIO:
    """Export report as Word document, returns BytesIO."""
    buf = BytesIO()
    doc = Document()
    if not report_text or not report_text.strip():
        doc.add_paragraph("Report content is empty.")
        doc.save(buf)
        buf.seek(0)
        return buf
    doc.add_heading(metadata.get("title", "EdgeRef AI Pre-submission Review Report"), level=0)

    doc.add_paragraph(f"Generated: {metadata.get('generated_at', '')}")
    doc.add_paragraph(f"Review Mode: {metadata.get('mode', '')}")
    doc.add_paragraph(f"Review Engine: {metadata.get('engine', '')}")
    if metadata.get("manuscript_language"):
        doc.add_paragraph(f"Manuscript Language: {metadata['manuscript_language']}")
    if metadata.get("review_output_language"):
        doc.add_paragraph(f"Review Output Language: {metadata['review_output_language']}")
    if metadata.get("journal_match_status"):
        doc.add_paragraph(f"Journal Catalog Match: {metadata['journal_match_status']}")
        if metadata.get("journal_match_confidence"):
            doc.add_paragraph(f"  Confidence: {metadata['journal_match_confidence']}")
        if metadata.get("journal_match_count"):
            doc.add_paragraph(f"  Matched records: {metadata['journal_match_count']}")
        if metadata.get("journal_match_journal"):
            doc.add_paragraph(f"  Matched Journal: {metadata['journal_match_journal']}")
        if metadata.get("journal_match_database"):
            doc.add_paragraph(f"  Database: {metadata['journal_match_database']}")
        if metadata.get("journal_match_db_type"):
            doc.add_paragraph(f"  DB Type: {metadata['journal_match_db_type']}")
        if metadata.get("journal_match_version"):
            doc.add_paragraph(f"  Version: {metadata['journal_match_version']}")
        if metadata.get("journal_match_issn"):
            doc.add_paragraph(f"  ISSN: {metadata['journal_match_issn']}")
        if metadata.get("journal_match_cn"):
            doc.add_paragraph(f"  CN: {metadata['journal_match_cn']}")

    if metadata.get("target_journal"):
        doc.add_paragraph(f"Target Journal: {metadata['target_journal']}")
    if metadata.get("subject_area"):
        doc.add_paragraph(f"Subject Area: {metadata['subject_area']}")
    if metadata.get("article_type"):
        doc.add_paragraph(f"Article Type: {metadata['article_type']}")

    doc.add_paragraph("")

    for raw_line in (report_text or "").splitlines():
        line = raw_line.strip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("- [ ] "):
            doc.add_paragraph(line[6:].strip(), style="List Bullet")
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith("1. ") or line.startswith("2. ") or line.startswith("3. ") or line.startswith("4. ") or line.startswith("5. ") or line.startswith("6. ") or line.startswith("7. ") or line.startswith("8. ") or line.startswith("9. "):
            doc.add_paragraph(line, style="List Number")
        else:
            doc.add_paragraph(raw_line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
